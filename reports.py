from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_trial_balance_pdf(data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Title'], alignment=TA_CENTER, fontSize=18)
    elements.append(Paragraph(f"Trial Balance - {data['fiscal_year']}", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Table data
    table_data = [['Account Code', 'Account Name (AR)', 'Account Name (EN)', 'Debit', 'Credit', 'Balance']]
    
    for item in data['trial_balance']:
        table_data.append([
            item['account_code'],
            item['account_name_ar'],
            item['account_name_en'],
            f"{item['debit']:.2f}",
            f"{item['credit']:.2f}",
            f"{item['balance']:.2f}"
        ])
    
    table_data.append(['', '', 'TOTAL', f"{data['total_debit']:.2f}", f"{data['total_credit']:.2f}", ''])
    
    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_trial_balance_excel(data):
    buffer = BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Trial Balance"
    
    # Title
    ws.merge_cells('A1:F1')
    ws['A1'] = f"Trial Balance - {data['fiscal_year']}"
    ws['A1'].font = Font(size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # Headers
    headers = ['Account Code', 'Account Name (AR)', 'Account Name (EN)', 'Debit', 'Credit', 'Balance']
    ws.append([])
    ws.append(headers)
    header_row = ws[3]
    for cell in header_row:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')
    
    # Data
    for item in data['trial_balance']:
        ws.append([
            item['account_code'],
            item['account_name_ar'],
            item['account_name_en'],
            item['debit'],
            item['credit'],
            item['balance']
        ])
    
    # Total
    ws.append(['', '', 'TOTAL', data['total_debit'], data['total_credit'], ''])
    total_row = ws[ws.max_row]
    for cell in total_row:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
    
    # Adjust column widths
    from openpyxl.utils import get_column_letter
    for col_idx, column in enumerate(ws.columns, 1):
        max_length = 0
        column_letter = get_column_letter(col_idx)
        for cell in column:
            if cell.value and not isinstance(cell, type(None)):
                try:
                    max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
        ws.column_dimensions[column_letter].width = max_length + 2
    
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generate_trial_balance_word(data):
    buffer = BytesIO()
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Trial Balance - {data['fiscal_year']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ['Account Code', 'Account Name (AR)', 'Account Name (EN)', 'Debit', 'Credit', 'Balance']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data
    for item in data['trial_balance']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['account_code']
        row_cells[1].text = item['account_name_ar']
        row_cells[2].text = item['account_name_en']
        row_cells[3].text = f"{item['debit']:.2f}"
        row_cells[4].text = f"{item['credit']:.2f}"
        row_cells[5].text = f"{item['balance']:.2f}"
    
    # Total
    total_row = table.add_row().cells
    total_row[2].text = 'TOTAL'
    total_row[3].text = f"{data['total_debit']:.2f}"
    total_row[4].text = f"{data['total_credit']:.2f}"
    for cell in total_row:
        cell.paragraphs[0].runs[0].font.bold = True
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_account_statement_pdf(data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Title'], alignment=TA_CENTER, fontSize=18)
    elements.append(Paragraph(f"Account Statement - {data['account']['account_code']}", title_style))
    elements.append(Paragraph(f"{data['account']['account_name_ar']} / {data['account']['account_name_en']}", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Table data
    table_data = [['Date', 'Description', 'Debit', 'Credit', 'Balance']]
    
    for txn in data['transactions']:
        table_data.append([
            str(txn['date'])[:10],
            txn['description'],
            f"{txn['debit']:.2f}",
            f"{txn['credit']:.2f}",
            f"{txn['balance']:.2f}"
        ])
    
    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_account_statement_excel(data):
    buffer = BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Account Statement"
    
    # Title
    ws.merge_cells('A1:E1')
    ws['A1'] = f"Account Statement - {data['account']['account_code']}"
    ws['A1'].font = Font(size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    ws.merge_cells('A2:E2')
    ws['A2'] = f"{data['account']['account_name_ar']} / {data['account']['account_name_en']}"
    ws['A2'].font = Font(size=12)
    ws['A2'].alignment = Alignment(horizontal='center')
    
    # Headers
    headers = ['Date', 'Description', 'Debit', 'Credit', 'Balance']
    ws.append([])
    ws.append(headers)
    header_row = ws[4]
    for cell in header_row:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')
    
    # Data
    for txn in data['transactions']:
        ws.append([
            str(txn['date'])[:10],
            txn['description'],
            txn['debit'],
            txn['credit'],
            txn['balance']
        ])
    
    # Adjust column widths
    from openpyxl.utils import get_column_letter
    for col_idx, column in enumerate(ws.columns, 1):
        max_length = 0
        column_letter = get_column_letter(col_idx)
        for cell in column:
            if cell.value and not isinstance(cell, type(None)):
                try:
                    max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
        ws.column_dimensions[column_letter].width = max_length + 2
    
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generate_account_statement_word(data):
    buffer = BytesIO()
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Account Statement - {data['account']['account_code']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"{data['account']['account_name_ar']} / {data['account']['account_name_en']}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ['Date', 'Description', 'Debit', 'Credit', 'Balance']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data
    for txn in data['transactions']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(txn['date'])[:10]
        row_cells[1].text = txn['description']
        row_cells[2].text = f"{txn['debit']:.2f}"
        row_cells[3].text = f"{txn['credit']:.2f}"
        row_cells[4].text = f"{txn['balance']:.2f}"
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_general_ledger_pdf(data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Title'], alignment=TA_CENTER, fontSize=18)
    elements.append(Paragraph(f"General Ledger - {data['fiscal_year']}", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    for account_data in data['ledger']:
        account = account_data['account']
        
        # Account header
        elements.append(Paragraph(f"<b>{account['account_code']} - {account['account_name_ar']} / {account['account_name_en']}</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.2*cm))
        
        # Table data
        table_data = [['Date', 'Description', 'Debit', 'Credit', 'Balance']]
        
        for txn in account_data['transactions']:
            table_data.append([
                str(txn['date'])[:10],
                txn['description'][:50],
                f"{txn['debit']:.2f}",
                f"{txn['credit']:.2f}",
                f"{txn['balance']:.2f}"
            ])
        
        if len(table_data) > 1:
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(table)
        
        elements.append(Spacer(1, 0.5*cm))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_general_ledger_excel(data):
    buffer = BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "General Ledger"
    
    row_num = 1
    
    # Title
    ws.merge_cells(f'A{row_num}:E{row_num}')
    ws[f'A{row_num}'] = f"General Ledger - {data['fiscal_year']}"
    ws[f'A{row_num}'].font = Font(size=16, bold=True)
    ws[f'A{row_num}'].alignment = Alignment(horizontal='center')
    row_num += 2
    
    for account_data in data['ledger']:
        account = account_data['account']
        
        # Account header
        ws.merge_cells(f'A{row_num}:E{row_num}')
        ws[f'A{row_num}'] = f"{account['account_code']} - {account['account_name_ar']} / {account['account_name_en']}"
        ws[f'A{row_num}'].font = Font(size=12, bold=True)
        row_num += 1
        
        # Headers
        headers = ['Date', 'Description', 'Debit', 'Credit', 'Balance']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row_num, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        row_num += 1
        
        # Data
        for txn in account_data['transactions']:
            ws.cell(row=row_num, column=1, value=str(txn['date'])[:10])
            ws.cell(row=row_num, column=2, value=txn['description'])
            ws.cell(row=row_num, column=3, value=txn['debit'])
            ws.cell(row=row_num, column=4, value=txn['credit'])
            ws.cell(row=row_num, column=5, value=txn['balance'])
            row_num += 1
        
        row_num += 1
    
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generate_general_ledger_word(data):
    buffer = BytesIO()
    doc = Document()
    
    # Title
    title = doc.add_heading(f"General Ledger - {data['fiscal_year']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for account_data in data['ledger']:
        account = account_data['account']
        
        # Account header
        doc.add_heading(f"{account['account_code']} - {account['account_name_ar']} / {account['account_name_en']}", level=2)
        
        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ['Date', 'Description', 'Debit', 'Credit', 'Balance']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data
        for txn in account_data['transactions']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(txn['date'])[:10]
            row_cells[1].text = txn['description']
            row_cells[2].text = f"{txn['debit']:.2f}"
            row_cells[3].text = f"{txn['credit']:.2f}"
            row_cells[4].text = f"{txn['balance']:.2f}"
        
        doc.add_paragraph()  # Spacing
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer